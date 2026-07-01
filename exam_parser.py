import csv
import json
import os
import re
from collections import Counter

try:
    import fitz
except ImportError:
    fitz = None

try:
    import pandas as pd
except ImportError:
    pd = None


ARABIC_DIGITS = str.maketrans(
    "\u0660\u0661\u0662\u0663\u0664\u0665\u0666\u0667\u0668\u0669\u06f0\u06f1\u06f2\u06f3\u06f4\u06f5\u06f6\u06f7\u06f8\u06f9",
    "01234567890123456789",
)

AR_QUESTION_WORD = r"(?:\u0627\u0644)?\u0633\u0624\u0627\u0644|\u0633\u0624\u0627\u0644|\u0633"
AR_ORDINALS = {
    "\u0627\u0644\u0623\u0648\u0644": 1,
    "\u0627\u0644\u0627\u0648\u0644": 1,
    "\u0623\u0648\u0644": 1,
    "\u0627\u0648\u0644": 1,
    "\u0627\u0644\u062b\u0627\u0646\u064a": 2,
    "\u062b\u0627\u0646\u064a": 2,
    "\u0627\u0644\u062b\u0627\u0644\u062b": 3,
    "\u062b\u0627\u0644\u062b": 3,
    "\u0627\u0644\u0631\u0627\u0628\u0639": 4,
    "\u0631\u0627\u0628\u0639": 4,
    "\u0627\u0644\u062e\u0627\u0645\u0633": 5,
    "\u062e\u0627\u0645\u0633": 5,
    "\u0627\u0644\u0633\u0627\u062f\u0633": 6,
    "\u0633\u0627\u062f\u0633": 6,
    "\u0627\u0644\u0633\u0627\u0628\u0639": 7,
    "\u0633\u0627\u0628\u0639": 7,
    "\u0627\u0644\u062b\u0627\u0645\u0646": 8,
    "\u062b\u0627\u0645\u0646": 8,
    "\u0627\u0644\u062a\u0627\u0633\u0639": 9,
    "\u062a\u0627\u0633\u0639": 9,
    "\u0627\u0644\u0639\u0627\u0634\u0631": 10,
    "\u0639\u0627\u0634\u0631": 10,
}

AR_ORDINAL_PATTERN = "|".join(re.escape(key) for key in sorted(AR_ORDINALS, key=len, reverse=True))
NUMBER_TOKEN = r"[0-9\u0660-\u0669\u06f0-\u06f9]{1,3}"
QUESTION_START_RE = re.compile(
    rf"^\s*(?:"
    rf"(?:Q|Qu|Que|Ques|Question)\s*(?:No\.?|#|[-.:])?\s*({NUMBER_TOKEN})\b"
    rf"|{AR_QUESTION_WORD}\s*(?:\u0631\u0642\u0645|#|[-.:])?\s*(?:({NUMBER_TOKEN})|({AR_ORDINAL_PATTERN}))"
    rf"|({NUMBER_TOKEN})\s*[\).:\-\u2013\u2014]"
    rf")",
    re.IGNORECASE,
)

OPTION_RE = re.compile(
    r"(?ms)^\s*([A-Da-d\u0623\u0628\u062c\u062f\u0647\u0640])[\).:-]\s*(.*?)(?=^\s*[A-Da-d\u0623\u0628\u062c\u062f\u0647\u0640][\).:-]\s*|\Z)"
)
SUBQUESTION_RE = re.compile(r"^\s*(?:[a-z]|[ivx]{1,5}|[A-Z])\s*[\).]", re.IGNORECASE)
CLO_RE = re.compile(r"\bCLO\s*[-:]?\s*(\d+(?:\.\d+)*)\b|\b(?:LO|L\.O\.)\s*[-:]?\s*(\d+(?:\.\d+)*)\b", re.IGNORECASE)

IMPERATIVE_RE = re.compile(
    r"\b(explain|calculate|define|discuss|compare|analyze|analyse|evaluate|choose|select|identify|write|solve|derive|design|implement|trace|prove|describe)\b"
    r"|(?:\u0627\u062e\u062a\u0631|\u0627\u0634\u0631\u062d|\u0627\u062d\u0633\u0628|\u0639\u0631\u0641|\u0646\u0627\u0642\u0634|\u0642\u0627\u0631\u0646|\u062d\u0644\u0644|\u0642\u064a\u0645|\u062d\u062f\u062f|\u0627\u0643\u062a\u0628|\u0635\u0645\u0645)",
    re.IGNORECASE,
)

MARKS_RE = re.compile(
    r"[\[\(]?\s*(\d+(?:\.\d+)?)\s*(?:marks?|points?|\u062f\u0631\u062c\u0627\u062a?|\u062f\u0631\u062c\u0629)\s*[\]\)]?"
    r"|(?:marks?|points?|\u062f\u0631\u062c\u0627\u062a?|\u062f\u0631\u062c\u0629)\s*[:=]\s*(\d+(?:\.\d+)?)",
    re.IGNORECASE,
)

ANSWER_KEY_RE = re.compile(
    r"\b(answer\s+key|answer\s+sheet|answers?\s+table|model\s+answer|marking\s+scheme|solution\s+key|clo\s+coverage\s+summary|coverage\s+summary)\b"
    r"|(?:\u0645\u0641\u062a\u0627\u062d\s+\u0627\u0644\u0625\u062c\u0627\u0628|\u0646\u0645\u0648\u0630\u062c\s+\u0627\u0644\u0625\u062c\u0627\u0628|\u062c\u062f\u0648\u0644\s+\u0627\u0644\u0625\u062c\u0627\u0628|\u062a\u063a\u0637\u064a\u0629\s+.*\u0646\u0648\u0627\u062a\u062c)",
    re.IGNORECASE,
)

MATH_SYMBOL_REPLACEMENTS = {
    "\u00d7": " x ",
    "\u2212": "-",
}


class ExamParser:
    def __init__(self, filepath):
        self.filepath = filepath
        self.file_ext = os.path.splitext(filepath)[1].lower()
        self.raw_text = ""
        self.pages = []
        self.document_blocks = []
        self.normalized_blocks = []
        self.question_candidates = []
        self.question_segments = []
        self.metadata = {
            "course_code": None,
            "exam_type": None,
            "total_marks": None,
            "course_name": None,
        }
        self.parsed_questions = []
        self.validation_warnings = []

    def parse(self):
        self.text_extraction_layer()
        self.normalization_layer()
        self.extract_metadata()
        self.question_candidates = self.candidate_detection_layer()
        self.question_segments = self.segmentation_layer(self.question_candidates)
        self.parsed_questions = [self.extraction_layer(segment) for segment in self.question_segments]
        self.validation_layer()
        return self.parsed_questions

    # Backward-compatible entry point used by existing code.
    def extract_text(self):
        self.text_extraction_layer()
        return self.raw_text

    def text_extraction_layer(self):
        self.raw_text = ""
        self.pages = []
        self.document_blocks = []

        if self.file_ext == ".pdf":
            self._extract_pdf_blocks()
        elif self.file_ext == ".docx":
            self._extract_docx_blocks()
        elif self.file_ext == ".txt":
            self._extract_txt_lines()
        else:
            raise ValueError("Unsupported file format. Please use PDF, DOCX, or TXT.")

        self.document_blocks.sort(key=lambda item: (item["page"], item["order"]))
        self.pages = self._pages_from_blocks(self.document_blocks)
        self.raw_text = "\n".join(block["text"] for block in self.document_blocks if block.get("text")).strip()

    def _extract_pdf_blocks(self):
        if fitz is None:
            raise ImportError("PyMuPDF is required for PDF parsing.")
        doc = fitz.open(self.filepath)
        order = 0
        for page_index in range(len(doc)):
            page = doc.load_page(page_index)
            page_dict = page.get_text("dict")
            for block in page_dict.get("blocks", []):
                if block.get("type") != 0:
                    continue
                lines = []
                font_sizes = []
                font_flags = []
                line_ys = []
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    line_text = "".join(span.get("text", "") for span in spans).strip()
                    if not line_text:
                        continue
                    lines.append(line_text)
                    if line.get("bbox"):
                        line_ys.append(line["bbox"][1])
                    for span in spans:
                        if span.get("size"):
                            font_sizes.append(float(span["size"]))
                        font_flags.append(int(span.get("flags", 0)))
                text = "\n".join(lines).strip()
                if not text:
                    continue
                line_spacing = None
                if len(line_ys) > 1:
                    diffs = [line_ys[i + 1] - line_ys[i] for i in range(len(line_ys) - 1)]
                    line_spacing = round(sum(diffs) / len(diffs), 2)
                self.document_blocks.append({
                    "text": text,
                    "page": page_index + 1,
                    "order": order,
                    "source": "pdf",
                    "block_type": "text",
                    "font_size": round(sum(font_sizes) / len(font_sizes), 2) if font_sizes else None,
                    "bold": any(flag & 16 for flag in font_flags),
                    "bbox": tuple(round(v, 2) for v in block.get("bbox", (0, 0, 0, 0))),
                    "line_spacing": line_spacing,
                    "style": "",
                })
                order += 1

    def _extract_docx_blocks(self):
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required for DOCX parsing.")

        document = Document(self.filepath)
        order = 0
        numbering_index = 1
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style = paragraph.style.name if paragraph.style else ""
            numbered = "list" in style.lower()
            prefix = f"{numbering_index}. " if numbered and not QUESTION_START_RE.search(text) else ""
            if numbered:
                numbering_index += 1
            bold_runs = [run.bold for run in paragraph.runs if run.text.strip()]
            font_sizes = [
                run.font.size.pt
                for run in paragraph.runs
                if run.text.strip() and run.font.size is not None
            ]
            self.document_blocks.append({
                "text": prefix + text,
                "page": 1,
                "order": order,
                "source": "docx",
                "block_type": "paragraph",
                "font_size": round(sum(font_sizes) / len(font_sizes), 2) if font_sizes else None,
                "bold": bool(bold_runs and all(value for value in bold_runs)),
                "bbox": None,
                "line_spacing": None,
                "style": style,
                "numbered": numbered,
            })
            order += 1

        for table in document.tables:
            rows = []
            for row in table.rows:
                cells = [normalize_whitespace(cell.text) for cell in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            if not rows:
                continue
            self.document_blocks.append({
                "text": "\n".join(rows),
                "page": 1,
                "order": order,
                "source": "docx",
                "block_type": "table",
                "font_size": None,
                "bold": False,
                "bbox": None,
                "line_spacing": None,
                "style": "table",
            })
            order += 1

    def _extract_txt_lines(self):
        with open(self.filepath, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        for order, line in enumerate(text.splitlines()):
            if line.strip():
                self.document_blocks.append({
                    "text": line.rstrip(),
                    "page": 1,
                    "order": order,
                    "source": "txt",
                    "block_type": "line",
                    "font_size": None,
                    "bold": False,
                    "bbox": None,
                    "line_spacing": None,
                    "style": "",
                })

    def _pages_from_blocks(self, blocks):
        pages = {}
        for block in blocks:
            pages.setdefault(block["page"], []).append(block["text"])
        return [{"page": page, "text": "\n".join(lines)} for page, lines in sorted(pages.items())]

    def normalization_layer(self):
        blocks = []
        line_counter = Counter()
        page_count = max((block["page"] for block in self.document_blocks), default=1)
        for block in self.document_blocks:
            for line in str(block.get("text") or "").splitlines():
                normalized = normalize_for_repetition(line)
                if normalized:
                    line_counter[normalized] += 1

        repeated = {
            line for line, count in line_counter.items()
            if page_count > 1 and count >= max(2, int(page_count * 0.6)) and not QUESTION_START_RE.search(line)
        }

        for block in self.document_blocks:
            cleaned_lines = []
            for line in str(block.get("text") or "").splitlines():
                normalized_for_repeat = normalize_for_repetition(line)
                if normalized_for_repeat in repeated:
                    continue
                cleaned = self.normalize_text(line)
                if cleaned:
                    cleaned_lines.append(cleaned)
            text = "\n".join(cleaned_lines).strip()
            if not text:
                continue
            item = dict(block)
            item["text"] = text
            item["normalized_text"] = text
            blocks.append(item)

        self.normalized_blocks = blocks
        self.raw_text = "\n".join(block["text"] for block in blocks).strip()
        self.pages = self._pages_from_blocks(blocks)
        return blocks

    def normalize_text(self, value):
        text = str(value or "").translate(ARABIC_DIGITS)
        for old, new in MATH_SYMBOL_REPLACEMENTS.items():
            text = text.replace(old, new)
        text = re.sub(r"(?<=[A-Za-z\u0600-\u06ff])\s+(?=[,.;:!?؟])", "", text)
        text = re.sub(r"(?<=\d)\s+(?=[\+\-\*/=<>≤≥])", "", text)
        text = re.sub(r"(?<=[\+\-\*/=<>≤≥])\s+(?=\d)", "", text)
        text = re.sub(r"[ \t]{2,}", " ", text)
        text = re.sub(r"\s+([)\]])", r"\1", text)
        text = re.sub(r"([(\[])\s+", r"\1", text)
        return text.strip()

    def candidate_detection_layer(self):
        candidates = []
        body_blocks = self._blocks_before_answer_key(self.normalized_blocks)
        font_sizes = [block["font_size"] for block in body_blocks if block.get("font_size")]
        avg_font = sum(font_sizes) / len(font_sizes) if font_sizes else None
        current_section = ""

        for index, block in enumerate(body_blocks):
            text = block["text"]
            if self._looks_like_section_heading(block):
                current_section = text
            first_line = text.splitlines()[0] if text else ""
            evidence = []
            score = 0.0
            number = None

            match = QUESTION_START_RE.search(first_line)
            if match:
                number_text = next((group for group in match.groups() if group), None)
                number = self.normalize_question_number(number_text, len(candidates) + 1)
                evidence.append("numbering_pattern")
                score += 0.48

            if MARKS_RE.search(text):
                evidence.append("marks_pattern")
                score += 0.16
            if IMPERATIVE_RE.search(text):
                evidence.append("imperative_verb")
                score += 0.13
            if block.get("bold"):
                evidence.append("bold_text")
                score += 0.08
            if avg_font and block.get("font_size") and block["font_size"] >= avg_font + 1:
                evidence.append("larger_font")
                score += 0.08
            if block.get("bbox") and block["bbox"][0] < 120:
                evidence.append("left_margin_position")
                score += 0.04
            if current_section:
                evidence.append(f"section:{current_section[:60]}")
                score += 0.03

            if self._looks_like_option_line(first_line) or self._looks_like_answer_key_row(text):
                score -= 0.45
                evidence.append("suppressed_option_or_answer_key")
            if SUBQUESTION_RE.search(first_line) and not match:
                score -= 0.18
                evidence.append("possible_subquestion")

            if score >= 0.35:
                candidates.append({
                    "block_index": index,
                    "number": number,
                    "confidence": round(min(score, 1.0), 2),
                    "evidence": evidence,
                    "section": current_section,
                    "page": block["page"],
                    "text": text,
                })

        if not candidates:
            return self._heuristic_candidates(body_blocks)
        return candidates

    def _heuristic_candidates(self, blocks):
        candidates = []
        for index, block in enumerate(blocks):
            text = block.get("text", "")
            if IMPERATIVE_RE.search(text) or MARKS_RE.search(text) or self.extract_mcq_options(text):
                candidates.append({
                    "block_index": index,
                    "number": len(candidates) + 1,
                    "confidence": 0.42,
                    "evidence": ["heuristic_question_signal"],
                    "section": "",
                    "page": block["page"],
                    "text": text,
                })
        if candidates:
            return candidates
        if blocks:
            return [{
                "block_index": 0,
                "number": 1,
                "confidence": 0.25,
                "evidence": ["fallback_single_block"],
                "section": "",
                "page": blocks[0]["page"],
                "text": blocks[0]["text"],
            }]
        return []

    def segmentation_layer(self, candidates):
        blocks = self._blocks_before_answer_key(self.normalized_blocks)
        if not candidates:
            return []

        accepted = []
        seen_numbers = set()
        for candidate in sorted(candidates, key=lambda item: item["block_index"]):
            block = blocks[candidate["block_index"]]
            first_line = block["text"].splitlines()[0] if block.get("text") else ""
            if self._looks_like_option_line(first_line):
                continue
            number = candidate.get("number")
            if number is not None and number in seen_numbers:
                continue
            if number is not None:
                seen_numbers.add(number)
            accepted.append(candidate)

        segments = []
        for position, candidate in enumerate(accepted):
            start = candidate["block_index"]
            end = accepted[position + 1]["block_index"] if position + 1 < len(accepted) else len(blocks)
            segment_blocks = blocks[start:end]
            text = self._merge_question_blocks(segment_blocks)
            text = self.clean_question_block(text)
            if not text or self._looks_like_answer_key_row(text):
                continue
            page_start = segment_blocks[0]["page"]
            page_end = segment_blocks[-1]["page"]
            segments.append({
                "number": candidate.get("number") or len(segments) + 1,
                "section": candidate.get("section") or "",
                "text": text,
                "blocks": segment_blocks,
                "page_start": page_start,
                "page_end": page_end,
                "candidate_confidence": candidate.get("confidence", 0.5),
                "candidate_evidence": candidate.get("evidence", []),
            })
        return segments

    def _merge_question_blocks(self, blocks):
        parts = []
        for block in blocks:
            text = block.get("text", "").strip()
            if not text:
                continue
            if block.get("block_type") == "table":
                parts.append(text)
                continue
            if parts and self._line_continues_previous(parts[-1], text):
                parts[-1] = parts[-1].rstrip() + " " + text.lstrip()
            else:
                parts.append(text)
        return "\n".join(parts)

    def _line_continues_previous(self, previous, current):
        previous = str(previous or "").rstrip()
        current = str(current or "").lstrip()
        if not previous or not current:
            return False
        if self._looks_like_option_line(current) or QUESTION_START_RE.search(current):
            return False
        if previous.endswith((".", "?", "؟", ":", ";")):
            return False
        return len(previous) < 180 or current[:1].islower()

    def classification_layer(self, segment):
        text = segment.get("text", "")
        lowered = text.lower()
        reasons = []
        q_type = "unknown"
        confidence = 0.35

        if self.extract_mcq_options(text):
            q_type = "multiple_choice"
            confidence = 0.9
            reasons.append("detected_answer_choices")
        elif re.search(r"\b(true|false)\b|(?:\u0635\u062d|\u062e\u0637\u0623|\u0635\u0648\u0627\u0628|\u063a\u0644\u0637)", lowered, re.IGNORECASE):
            q_type = "true_false"
            confidence = 0.85
            reasons.append("true_false_terms")
        elif "___" in text or re.search(r"\b(fill|complete)\b|(?:\u0623\u0643\u0645\u0644|\u0627\u0645\u0644\u0623)", lowered):
            q_type = "fill_blank"
            confidence = 0.82
            reasons.append("blank_or_complete_terms")
        elif re.search(r"\b(match|matching|connect)\b|(?:\u0637\u0627\u0628\u0642|\u0648\u0635\u0644)", lowered):
            q_type = "matching"
            confidence = 0.76
            reasons.append("matching_terms")
        elif re.search(r"\b(calculate|solve|compute|derive)\b|(?:\u0627\u062d\u0633\u0628|\u062d\u0644|\u0627\u0634\u062a\u0642)", lowered) or re.search(r"[=+\-*/≤≥]", text):
            q_type = "calculation"
            confidence = 0.72
            reasons.append("calculation_terms_or_symbols")
        elif re.search(r"\b(code|program|python|java|sql|algorithm|function)\b|(?:\u0628\u0631\u0645\u062c|\u062e\u0648\u0627\u0631\u0632\u0645)", lowered):
            q_type = "programming"
            confidence = 0.72
            reasons.append("programming_terms")
        elif re.search(r"\b(case|scenario|patient|study)\b|(?:\u062d\u0627\u0644\u0629|\u0633\u064a\u0646\u0627\u0631\u064a\u0648)", lowered):
            q_type = "case_study"
            confidence = 0.68
            reasons.append("case_study_context")
        elif re.search(r"\b(define|list|identify|state|name)\b|(?:\u0639\u0631\u0641|\u0639\u062f\u062f|\u062d\u062f\u062f|\u0627\u0630\u0643\u0631)", lowered):
            q_type = "short_answer"
            confidence = 0.68
            reasons.append("short_answer_imperative")
        elif re.search(r"\b(explain|discuss|compare|analyze|analyse|evaluate|describe)\b|(?:\u0627\u0634\u0631\u062d|\u0646\u0627\u0642\u0634|\u0642\u0627\u0631\u0646|\u062d\u0644\u0644|\u0642\u064a\u0645|\u0635\u0641)", lowered):
            q_type = "essay"
            confidence = 0.7
            reasons.append("essay_imperative")

        return {
            "question_type": q_type,
            "confidence": round(confidence, 2),
            "reasons": reasons or ["no_strong_type_signal"],
        }

    def extraction_layer(self, segment):
        classification = self.classification_layer(segment)
        text = segment["text"].strip()
        marks = self.extract_marks(text)
        choices = self.extract_mcq_options(text)
        subquestions = self.extract_subquestions(text)
        clo_tags = self.detect_clo_tags(text)
        confidence = min(
            1.0,
            (segment.get("candidate_confidence", 0.5) * 0.65) + (classification["confidence"] * 0.35),
        )
        warnings = self.segment_warnings(segment, text, choices)
        number = segment.get("number") or len(self.parsed_questions) + 1
        question_label = f"Q{number}"
        return {
            "question_id": question_label,
            "question_number": number,
            "question_label": question_label,
            "section": segment.get("section", ""),
            "question_text": text,
            "question_text_with_number": f"{question_label}: {text}",
            "question_type": classification["question_type"],
            "classification": classification,
            "options": choices,
            "answer_choices": choices,
            "marks": marks,
            "subquestions": subquestions,
            "detected_clo_tags": clo_tags,
            "mapped_clo": None,
            "page_number": segment.get("page_start", 1),
            "page_start": segment.get("page_start", 1),
            "page_end": segment.get("page_end", segment.get("page_start", 1)),
            "confidence": round(confidence, 2),
            "candidate_evidence": segment.get("candidate_evidence", []),
            "warnings": warnings,
        }

    def validation_layer(self):
        warnings = []
        numbers = []
        seen = set()

        for question in self.parsed_questions:
            number = question.get("question_number")
            if number is None:
                warnings.append("Missing question number.")
            else:
                if number in seen:
                    warnings.append(f"Duplicated question number: Q{number}.")
                seen.add(number)
                numbers.append(number)

            text = question.get("question_text", "")
            if not text.strip():
                warnings.append(f"{question.get('question_id', 'Question')} has no text.")
            if len(text.split()) < 3:
                warnings.append(f"{question.get('question_id', 'Question')} is suspiciously short.")
            if question.get("warnings"):
                warnings.extend(f"{question.get('question_id')}: {warning}" for warning in question["warnings"])

        if numbers:
            expected = list(range(min(numbers), max(numbers) + 1))
            if numbers != expected:
                warnings.append("Question numbers are not strictly sequential.")
            jumps = [
                (numbers[i], numbers[i + 1])
                for i in range(len(numbers) - 1)
                if numbers[i + 1] - numbers[i] > 1
            ]
            for before, after in jumps:
                warnings.append(f"Sudden numbering jump from Q{before} to Q{after}.")

        if self.metadata["total_marks"]:
            total = sum(float(question.get("marks") or 0) for question in self.parsed_questions)
            if abs(total - self.metadata["total_marks"]) > 0.5:
                warnings.append(
                    f"Sum of individual marks ({total}) does not equal total declared marks ({self.metadata['total_marks']})."
                )

        if ANSWER_KEY_RE.search(self.raw_text or ""):
            warnings.append("Possible answer key or solution section was detected and excluded from segmentation.")

        self.validation_warnings = sorted(dict.fromkeys(warnings))
        return self.validation_warnings

    # Backward-compatible method names.
    def detect_question_boundaries(self):
        if not self.normalized_blocks:
            self.normalization_layer()
        return [
            {
                "number": segment.get("number"),
                "text": segment.get("text"),
                "page": segment.get("page_start"),
                "page_end": segment.get("page_end"),
                "confidence": segment.get("candidate_confidence"),
                "evidence": segment.get("candidate_evidence", []),
            }
            for segment in self.segmentation_layer(self.candidate_detection_layer())
        ]

    def heuristic_question_detection(self):
        return self._heuristic_candidates(self._blocks_before_answer_key(self.normalized_blocks))

    def classify_question_type(self, text):
        q_type = self.classification_layer({"text": text}).get("question_type", "unknown")
        legacy = {
            "multiple_choice": "MCQ",
            "true_false": "T/F",
            "fill_blank": "Fill in the blank",
            "essay": "Essay",
            "short_answer": "Short Answer",
            "matching": "Matching",
            "calculation": "Calculation",
            "programming": "Programming",
            "case_study": "Case Study",
        }
        return legacy.get(q_type, "Unknown")

    def question_start_pattern(self):
        return QUESTION_START_RE

    def normalize_question_number(self, value, fallback):
        if value is None:
            return fallback
        text = str(value).strip().translate(ARABIC_DIGITS)
        if text in AR_ORDINALS:
            return AR_ORDINALS[text]
        try:
            return int(text)
        except ValueError:
            return fallback

    def exam_body_text(self):
        return "\n".join(block["text"] for block in self._blocks_before_answer_key(self.normalized_blocks or self.document_blocks))

    def clean_question_block(self, block):
        lines = [line.rstrip() for line in str(block or "").splitlines()]
        while lines and not lines[0].strip():
            lines.pop(0)
        if lines:
            match = QUESTION_START_RE.search(lines[0])
            if match:
                lines[0] = lines[0][match.end():].lstrip(" .:-\u2013\u2014")
        text = "\n".join(lines).strip()
        return re.sub(r"\n{3,}", "\n\n", text)

    def extract_mcq_options(self, text):
        options = {}
        for match in OPTION_RE.finditer(str(text or "")):
            key = match.group(1).upper()
            value = match.group(2).strip()
            if value:
                options[key] = value
        return options

    def extract_marks(self, text):
        for match in MARKS_RE.finditer(str(text or "")):
            value = next((group for group in match.groups() if group), None)
            if value:
                try:
                    return float(value)
                except ValueError:
                    continue
        return 1.0

    def extract_subquestions(self, text):
        subquestions = []
        current = None
        for line in str(text or "").splitlines():
            if SUBQUESTION_RE.search(line) and not self._looks_like_option_line(line):
                if current:
                    subquestions.append(current)
                current = {"label": line.split()[0].rstrip(")."), "text": line}
            elif current:
                current["text"] += "\n" + line
        if current:
            subquestions.append(current)
        return subquestions

    def detect_clo_tags(self, text):
        tags = []
        for match in CLO_RE.finditer(str(text or "")):
            tag = next((group for group in match.groups() if group), None)
            if tag and tag not in tags:
                tags.append(tag)
        return tags

    def extract_metadata(self):
        text = self.raw_text or ""
        code = re.search(r"\b([A-Za-z]{2,5}\s?\d{3,4})\b", text)
        if code:
            self.metadata["course_code"] = code.group(1).upper().replace(" ", "")

        lowered = text.lower()
        exam_types = {
            "midterm": "Midterm",
            "final": "Final",
            "quiz": "Quiz",
            "assignment": "Assignment",
            "\u0646\u0647\u0627\u0626\u064a": "Final",
            "\u0646\u0635\u0641\u064a": "Midterm",
            "\u0627\u062e\u062a\u0628\u0627\u0631 \u0642\u0635\u064a\u0631": "Quiz",
            "\u0642\u0635\u064a\u0631": "Quiz",
        }
        for cue, label in exam_types.items():
            if cue in lowered:
                self.metadata["exam_type"] = label
                break

        total = re.search(
            r"(?:total\s+marks?|score|\u0627\u0644\u062f\u0631\u062c\u0629\s+\u0627\u0644\u0643\u0644\u064a\u0629|\u0627\u0644\u0645\u062c\u0645\u0648\u0639)\s*[:=]?\s*(\d+(?:\.\d+)?)",
            text,
            re.IGNORECASE,
        )
        if total:
            self.metadata["total_marks"] = float(total.group(1))

    def validate_questions(self):
        return self.validation_layer()

    def segment_warnings(self, segment, text, choices):
        warnings = []
        if len(text.split()) < 3:
            warnings.append("suspiciously short question")
        if choices and len(text.splitlines()) <= len(choices):
            warnings.append("answer choices may be missing question stem")
        if ANSWER_KEY_RE.search(text):
            warnings.append("possible answer key included")
        return warnings

    def _blocks_before_answer_key(self, blocks):
        selected = []
        first_question_seen = False
        for block in blocks:
            text = block.get("text", "")
            if QUESTION_START_RE.search(text.splitlines()[0] if text else ""):
                first_question_seen = True
            if first_question_seen and ANSWER_KEY_RE.search(text):
                break
            selected.append(block)
        return selected

    def _looks_like_option_line(self, text):
        return bool(re.match(r"^\s*[A-Da-d\u0623\u0628\u062c\u062f\u0647\u0640][\).:-]\s+", str(text or "")))

    def _looks_like_answer_key_row(self, text):
        compact = re.sub(r"\s+", " ", str(text or "")).strip()
        return bool(re.fullmatch(r"[A-Da-d\u0623\u0628\u062c\u062f\u0647\u0640]\s*(?:CLO\s*\d+(?:\.\d+)*)?", compact))

    def _looks_like_section_heading(self, block):
        text = normalize_whitespace(block.get("text", ""))
        if not text or len(text) > 90:
            return False
        if QUESTION_START_RE.search(text) or self._looks_like_option_line(text):
            return False
        heading_terms = [
            "section", "part", "case study", "scenario",
            "\u0627\u0644\u0642\u0633\u0645", "\u0627\u0644\u062c\u0632\u0621", "\u062d\u0627\u0644\u0629",
        ]
        lowered = text.lower()
        return block.get("bold") or any(term in lowered for term in heading_terms)

    def export(self, export_format="json", output_path=None):
        data = {
            "exam_id": f"{self.metadata['course_code']}_{self.metadata['exam_type']}" if self.metadata["course_code"] else "Unknown_Exam",
            "metadata": self.metadata,
            "warnings": self.validation_warnings,
            "questions": self.parsed_questions,
            "candidates": self.question_candidates,
        }
        if export_format == "json":
            result = json.dumps(data, indent=2, ensure_ascii=False)
            if output_path:
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(result)
            return result
        if export_format == "csv":
            rows = self.export_rows()
            if output_path:
                with open(output_path, "w", encoding="utf-8-sig", newline="") as f:
                    writer = csv.DictWriter(f, fieldnames=list(rows[0].keys()) if rows else [])
                    writer.writeheader()
                    writer.writerows(rows)
            return rows
        if export_format == "excel":
            if pd is None:
                raise ImportError("pandas is required for Excel export.")
            df = pd.DataFrame(self.export_rows())
            if output_path:
                df.to_excel(output_path, index=False)
            return df
        return data

    def export_rows(self):
        return [
            {
                "Question ID": q["question_id"],
                "Type": q["question_type"],
                "Text": q["question_text"],
                "Marks": q["marks"],
                "Options": json.dumps(q["options"], ensure_ascii=False) if q["options"] else "",
                "Page": q["page_number"],
                "Confidence": q.get("confidence", ""),
                "Warnings": "; ".join(q.get("warnings", [])),
            }
            for q in self.parsed_questions
        ]


def normalize_whitespace(value):
    return re.sub(r"\s+", " ", str(value or "")).strip()


def normalize_for_repetition(value):
    text = normalize_whitespace(value).translate(ARABIC_DIGITS).lower()
    text = re.sub(r"\d+", "#", text)
    return text


if __name__ == "__main__":
    import sys
    sys.stdout.reconfigure(encoding="utf-8")
    if len(sys.argv) > 1:
        parser = ExamParser(sys.argv[1])
        parser.parse()
        print(parser.export(export_format="json"))
        if parser.validation_warnings:
            print("\nValidation Warnings:")
            for warning in parser.validation_warnings:
                print("-", warning)
    else:
        print("Usage: python exam_parser.py <path_to_pdf_or_docx_or_txt>")
